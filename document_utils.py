"""
Document utilities for Brag Document MCP.

This module contains helper functions for managing brag documents,
including path resolution, index file management, and document operations.
"""

import os
import json
import uuid
from pathlib import Path
from typing import Optional, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Using python-docx-ng for improved style handling

# Constants
DEFAULT_WORKSPACE_ROOT = os.getcwd()
TEMPLATE_PATH = Path("Templates") / "A brag document template.docx"
BRAG_DOCUMENTS_DIR = "BragDocuments"


def get_workspace_root(workspace_root: Optional[str] = None) -> Path:
    """Get the workspace root path.
    
    Args:
        workspace_root: Optional custom workspace root. If None, uses default.
    
    Returns:
        Path object representing the workspace root.
    """
    return Path(workspace_root) if workspace_root else Path(DEFAULT_WORKSPACE_ROOT)


def get_document_path(
    full_name: str,
    year: int,
    workspace_root: Optional[str] = None
) -> Path:
    """Get the path to the brag document.
    
    Args:
        full_name: Full name of the person (e.g., "John Doe")
        year: Year for the brag document (e.g., 2024)
        workspace_root: Optional root directory for documents.
    
    Returns:
        Path to the brag document file.
    """
    root = get_workspace_root(workspace_root)
    return root / BRAG_DOCUMENTS_DIR / full_name / f"Brag Document - {full_name} ({year}).docx"


def get_index_path(
    full_name: str,
    year: int,
    workspace_root: Optional[str] = None
) -> Path:
    """Get the path to the index JSON file.
    
    Args:
        full_name: Full name of the person (e.g., "John Doe")
        year: Year for the brag document (e.g., 2024)
        workspace_root: Optional root directory for documents.
    
    Returns:
        Path to the index JSON file.
    """
    root = get_workspace_root(workspace_root)
    return (
        root
        / BRAG_DOCUMENTS_DIR
        / full_name
        / ".index"
        / f"Brag Document - {full_name} ({year}).json"
    )


def get_template_path(workspace_root: Optional[str] = None) -> Path:
    """Get the path to the template document.
    
    Args:
        workspace_root: Optional root directory for documents.
    
    Returns:
        Path to the template document.
    """
    root = get_workspace_root(workspace_root)
    return root / TEMPLATE_PATH


def ensure_index_file(index_path: Path) -> None:
    """Create index file if it doesn't exist.
    
    Initializes the index file with an empty structure for tracking
    document sections and entries.
    
    Args:
        index_path: Path where the index file should be created.
    
    Raises:
        OSError: If the file cannot be created or written.
    """
    index_path.parent.mkdir(parents=True, exist_ok=True)
    
    if not index_path.exists():
        # Initialize with empty structure
        index_data = {
            "document_name": index_path.stem,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "sections": {},
            "entries": {}
        }
        with open(index_path, 'w', encoding='utf-8') as f:
            json.dump(index_data, f, indent=2, ensure_ascii=False)


def generate_entry_id() -> str:
    """Generate a unique entry ID.
    
    Returns:
        A unique entry ID string.
    """
    return str(uuid.uuid4())


def load_index(index_path: Path) -> dict:
    """Load the index file.
    
    Args:
        index_path: Path to the index file.
    
    Returns:
        Dictionary containing index data.
    
    Raises:
        FileNotFoundError: If the index file doesn't exist.
        json.JSONDecodeError: If the index file is invalid JSON.
    """
    if not index_path.exists():
        ensure_index_file(index_path)
    
    with open(index_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_index(index_path: Path, index_data: dict) -> None:
    """Save the index file.
    
    Args:
        index_path: Path to the index file.
        index_data: Dictionary containing index data to save.
    
    Raises:
        OSError: If the file cannot be written.
    """
    index_data["updated_at"] = datetime.now().isoformat()
    with open(index_path, 'w', encoding='utf-8') as f:
        json.dump(index_data, f, indent=2, ensure_ascii=False)


def is_heading_paragraph(paragraph) -> bool:
    """Check if a paragraph is a heading.
    
    Args:
        paragraph: A paragraph object from python-docx.
    
    Returns:
        True if the paragraph appears to be a heading, False otherwise.
    """
    if not paragraph.text.strip():
        return False
    
    # Check if it's a heading style
    if paragraph.style.name.startswith('Heading'):
        return True
    
    # Check if all runs are bold (common heading pattern)
    if paragraph.runs:
        all_bold = all(run.bold for run in paragraph.runs if run.text.strip())
        if all_bold:
            return True
    
    # Check if it's a short line (headings are often short)
    # and has some formatting that suggests it's a heading
    text = paragraph.text.strip()
    if len(text) < 100 and (paragraph.runs and paragraph.runs[0].bold):
        return True
    
    return False


def find_section_paragraph(doc: Document, section_path: str) -> Optional[Tuple[int, int]]:
    """Find the paragraph range for a section in the document.
    
    Searches for a section heading in the document and returns the
    start and end paragraph indices for that section.
    
    Args:
        doc: The Document object to search.
        section_path: The section path (e.g., "Projects" or "Outside of work/Articles").
    
    Returns:
        Tuple of (start_paragraph_index, end_paragraph_index) or None if not found.
        The end index is the index of the next section heading, or the last paragraph.
    """
    # Split section path for nested sections
    section_parts = [part.strip() for part in section_path.split('/')]
    main_section = section_parts[0]
    nested_section = section_parts[1] if len(section_parts) > 1 else None
    
    section_start = None
    section_end = None
    
    # Known top-level sections (in order they appear)
    top_level_sections = [
        "Goals for this year",
        "Goals for next year",
        "Goals for next year (optional)",  # Alternative name in template
        "Projects",
        "Collaboration & mentorship",
        "Documentation",
        "What you learned",
        "Outside of work"
    ]
    
    # Find the main section
    for i, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        
        # Check for exact match or alternative names
        is_match = para_text == main_section
        if not is_match and main_section == "Goals for next year":
            # Handle "Goals for next year (optional)" variation
            is_match = para_text == "Goals for next year (optional)"
        
        if is_heading_paragraph(paragraph) and is_match:
            section_start = i + 1  # Start after the heading
            # Find the next top-level section
            for next_section in top_level_sections:
                if next_section == main_section:
                    continue
                # Look ahead for the next section
                for j in range(i + 1, len(doc.paragraphs)):
                    next_para = doc.paragraphs[j]
                    next_para_text = next_para.text.strip()
                    if is_heading_paragraph(next_para) and next_para_text == next_section:
                        section_end = j
                        break
                if section_end:
                    break
            break
    
    if section_start is None:
        return None
    
    if section_end is None:
        section_end = len(doc.paragraphs)
    
    # If nested section, find it within the main section
    if nested_section:
        nested_start = None
        nested_end = section_end
        
        for i in range(section_start, section_end):
            para = doc.paragraphs[i]
            para_text = para.text.strip()
            
            if is_heading_paragraph(para) and para_text == nested_section:
                nested_start = i + 1
                # Find next heading or end of main section
                for j in range(i + 1, section_end):
                    next_para = doc.paragraphs[j]
                    if is_heading_paragraph(next_para):
                        nested_end = j
                        break
                break
        
        if nested_start is None:
            return None
        
        return (nested_start, nested_end)
    
    return (section_start, section_end)


def _ensure_style_exists(doc: Document, style_name: str) -> None:
    """Ensure a style exists in the document, creating it if necessary.
    
    Args:
        doc: The Document object.
        style_name: Name of the style to ensure exists.
    """
    style_names = [s.name for s in doc.styles]
    
    if style_name in style_names:
        return  # Style already exists
    
    # Try to add as latent style first (for built-in styles)
    try:
        latent = doc.styles.latent_styles
        try:
            ls = latent[style_name]
            ls.hidden = False
            ls.quick_style = True
        except KeyError:
            ls = latent.add_latent_style(style_name)
            ls.hidden = False
            ls.quick_style = True
            ls.priority = 1
    except (AttributeError, KeyError):
        # If latent_styles API doesn't work, add style explicitly
        doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def add_entry_to_document(
    doc_path: Path,
    section_path: str,
    text: str,
    position: Optional[int] = None
) -> int:
    """Add an entry to a section in the document.
    
    Args:
        doc_path: Path to the DOCX document.
        section_path: Section path where to add the entry.
        text: Text content of the entry.
        position: Optional position index to insert at (0-based). If None, appends.
    
    Returns:
        The paragraph index where the entry was added.
    
    Raises:
        FileNotFoundError: If the document doesn't exist.
        ValueError: If the section is not found.
    """
    # Open document with automatic style error recovery
    try:
        doc = Document(str(doc_path))
    except KeyError as e:
        if "List Bullet" in str(e):
            # Fix missing List Bullet style using opc package
            try:
                from docx import opc
                from docx.oxml import parse_xml
                import xml.etree.ElementTree as ET
                
                package = opc.Package.open(str(doc_path))
                document_part = package.main_document_part
                
                # Get or create styles part
                try:
                    styles_part = document_part.part_related_by(
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles'
                    )
                    styles_root = parse_xml(styles_part.blob)
                except KeyError:
                    from docx.opc.part import Part
                    from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
                    
                    styles_xml = (
                        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:style w:type="paragraph" w:styleId="ListBullet">'
                        '<w:name w:val="List Bullet"/>'
                        '<w:basedOn w:val="Normal"/>'
                        '<w:qFormat/>'
                        '</w:style>'
                        '</w:styles>'
                    )
                    styles_part = Part.new(
                        package,
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles',
                        CT.WML_STYLES,
                        parse_xml(styles_xml)
                    )
                    document_part.relate_to(styles_part, RT.STYLES)
                    package.save(str(doc_path))
                    doc = Document(str(doc_path))
                else:
                    # Check if ListBullet exists
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    style_elements = styles_root.findall('.//w:style', ns)
                    has_list_bullet = any(
                        (elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
                         or elem.get('styleId', '')) == 'ListBullet'
                        for elem in style_elements
                    )
                    
                    if not has_list_bullet:
                        style_elem = ET.SubElement(styles_root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style')
                        style_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'paragraph')
                        style_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId', 'ListBullet')
                        
                        name_elem = ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                        name_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'List Bullet')
                        
                        based_on = ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}basedOn')
                        based_on.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'Normal')
                        
                        ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}qFormat')
                        
                        styles_part.blob = styles_root
                        package.save(str(doc_path))
                    
                    # Retry opening
                    doc = Document(str(doc_path))
            except Exception:
                raise ValueError(
                    f"Document has style issues. The 'List Bullet' style is referenced but not defined. "
                    f"Please fix the template file or open '{doc_path}' in Microsoft Word and apply the 'List Bullet' style."
                ) from e
        else:
            raise
    
    # Ensure List Bullet style exists (in case it's missing)
    _ensure_style_exists(doc, 'List Bullet')
    
    # Find the section
    section_range = find_section_paragraph(doc, section_path)
    if section_range is None:
        raise ValueError(f"Section '{section_path}' not found in document")
    
    start_idx, end_idx = section_range
    
    # Determine insertion position
    if position is None:
        # Append: find the last non-empty paragraph in the section, or insert at start
        insert_idx = start_idx
        for i in range(start_idx, end_idx):
            if doc.paragraphs[i].text.strip():
                insert_idx = i + 1
    else:
        # Insert at specific position
        insert_idx = min(start_idx + position, end_idx)
    
    # Ensure insert_idx is within valid range
    insert_idx = min(insert_idx, len(doc.paragraphs))
    
    # Create new paragraph with bullet formatting
    # Insert paragraph at the correct position within the section
    if insert_idx < len(doc.paragraphs):
        # Insert empty paragraph first to avoid style inheritance issues
        new_para = doc.paragraphs[insert_idx].insert_paragraph_before()
    else:
        # Append to end if beyond document length
        new_para = doc.add_paragraph()
        insert_idx = len(doc.paragraphs) - 1
    
    # Add bullet character and text as a run
    bullet_run = new_para.add_run('• ' + text)
    
    # Format as list item with indentation
    new_para.paragraph_format.left_indent = Pt(18)
    
    # Save the document
    doc.save(str(doc_path))
    
    return insert_idx


def update_entry_to_document(
    doc_path: Path,
    paragraph_index: int,
    new_text: str
) -> None:
    """Update an entry in the document at a specific paragraph index.
    
    Args:
        doc_path: Path to the DOCX document.
        paragraph_index: Index of the paragraph to update.
        new_text: New text content for the entry (will be prefixed with bullet).
    
    Raises:
        FileNotFoundError: If the document doesn't exist.
        IndexError: If the paragraph index is out of range.
        ValueError: If the paragraph cannot be updated.
    """
    # Open document with automatic style error recovery
    try:
        doc = Document(str(doc_path))
    except KeyError as e:
        if "List Bullet" in str(e):
            # Fix missing List Bullet style using opc package
            try:
                from docx import opc
                from docx.oxml import parse_xml
                import xml.etree.ElementTree as ET
                
                package = opc.Package.open(str(doc_path))
                document_part = package.main_document_part
                
                # Get or create styles part
                try:
                    styles_part = document_part.part_related_by(
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles'
                    )
                    styles_root = parse_xml(styles_part.blob)
                except KeyError:
                    from docx.opc.part import Part
                    from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
                    
                    styles_xml = (
                        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:style w:type="paragraph" w:styleId="ListBullet">'
                        '<w:name w:val="List Bullet"/>'
                        '<w:basedOn w:val="Normal"/>'
                        '<w:qFormat/>'
                        '</w:style>'
                        '</w:styles>'
                    )
                    styles_part = Part.new(
                        package,
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles',
                        CT.WML_STYLES,
                        parse_xml(styles_xml)
                    )
                    document_part.relate_to(styles_part, RT.STYLES)
                    package.save(str(doc_path))
                    doc = Document(str(doc_path))
                else:
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    style_elements = styles_root.findall('.//w:style', ns)
                    has_list_bullet = any(
                        (elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
                         or elem.get('styleId', '')) == 'ListBullet'
                        for elem in style_elements
                    )
                    
                    if not has_list_bullet:
                        style_elem = ET.SubElement(styles_root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style')
                        style_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'paragraph')
                        style_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId', 'ListBullet')
                        
                        name_elem = ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                        name_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'List Bullet')
                        
                        based_on = ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}basedOn')
                        based_on.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'Normal')
                        
                        ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}qFormat')
                        
                        styles_part.blob = styles_root
                        package.save(str(doc_path))
                    
                    doc = Document(str(doc_path))
            except Exception:
                raise ValueError(
                    f"Document has style issues. The 'List Bullet' style is referenced but not defined. "
                    f"Please fix the template file or open '{doc_path}' in Microsoft Word and apply the 'List Bullet' style."
                ) from e
        else:
            raise
    
    # Check paragraph index is valid
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} is out of range. Document has {len(doc.paragraphs)} paragraphs.")
    
    # Get the paragraph to update
    para = doc.paragraphs[paragraph_index]
    
    # Check if paragraph already has bullet formatting
    # If it starts with bullet, preserve it; otherwise add it
    current_text = para.text.strip()
    if current_text.startswith('•') or current_text.startswith('-'):
        # Remove existing bullet and whitespace
        current_text = current_text.lstrip('•-').strip()
    
    # Update paragraph with new text (with bullet)
    bullet_text = '• ' + new_text
    para.clear()
    para.add_run(bullet_text)
    
    # Ensure formatting is preserved
    para.paragraph_format.left_indent = Pt(18)
    
    # Save the document
    doc.save(str(doc_path))


def update_entry_to_index(
    index_path: Path,
    entry_id: str,
    new_text: str
) -> None:
    """Update an entry in the index file.
    
    Args:
        index_path: Path to the index file.
        entry_id: Unique identifier for the entry to update.
        new_text: New text content for the entry.
    
    Raises:
        KeyError: If the entry_id is not found in the index.
    """
    index_data = load_index(index_path)
    
    # Check if entry exists
    if entry_id not in index_data["entries"]:
        raise KeyError(f"Entry with ID '{entry_id}' not found in index.")
    
    # Update entry data
    entry_data = index_data["entries"][entry_id]
    entry_data["text"] = new_text
    entry_data["updated_at"] = datetime.now().isoformat()
    
    # Update index metadata
    index_data["updated_at"] = datetime.now().isoformat()
    
    save_index(index_path, index_data)


def initialize_document_from_template(
    doc_path: Path,
    full_name: str,
    year: int
) -> None:
    """Initialize a newly created document from template.
    
    Replaces placeholders in the document title and removes example entries.
    
    Args:
        doc_path: Path to the DOCX document.
        full_name: Full name to replace <Full Name> placeholder.
        year: Year to replace <Year> placeholder.
    """
    # Open document with error recovery for style issues
    try:
        doc = Document(str(doc_path))
    except KeyError as e:
        if "List Bullet" in str(e):
            # Fix style issue first
            try:
                from docx import opc
                from docx.oxml import parse_xml
                import xml.etree.ElementTree as ET
                
                package = opc.Package.open(str(doc_path))
                document_part = package.main_document_part
                
                try:
                    styles_part = document_part.part_related_by(
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles'
                    )
                    styles_root = parse_xml(styles_part.blob)
                except KeyError:
                    from docx.opc.part import Part
                    from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
                    
                    styles_xml = (
                        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:style w:type="paragraph" w:styleId="ListBullet">'
                        '<w:name w:val="List Bullet"/>'
                        '<w:basedOn w:val="Normal"/>'
                        '<w:qFormat/>'
                        '</w:style>'
                        '</w:styles>'
                    )
                    styles_part = Part.new(
                        package,
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles',
                        CT.WML_STYLES,
                        parse_xml(styles_xml)
                    )
                    document_part.relate_to(styles_part, RT.STYLES)
                    package.save(str(doc_path))
                    doc = Document(str(doc_path))
                else:
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    style_elements = styles_root.findall('.//w:style', ns)
                    has_list_bullet = any(
                        (elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
                         or elem.get('styleId', '')) == 'ListBullet'
                        for elem in style_elements
                    )
                    
                    if not has_list_bullet:
                        style_elem = ET.SubElement(styles_root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style')
                        style_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'paragraph')
                        style_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId', 'ListBullet')
                        
                        name_elem = ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                        name_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'List Bullet')
                        
                        based_on = ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}basedOn')
                        based_on.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'Normal')
                        
                        ET.SubElement(style_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}qFormat')
                        
                        styles_part.blob = styles_root
                        package.save(str(doc_path))
                    
                    doc = Document(str(doc_path))
            except Exception:
                # If style fix fails, try to open anyway or skip initialization
                try:
                    doc = Document(str(doc_path))
                except Exception:
                    # If we still can't open, skip initialization
                    return
    
    # Replace title placeholders
    # Title should match the filename format: "Brag Document - <Full Name> (<Year>)"
    # Find and update the title paragraph (usually first paragraph)
    # The template may use em dash (–) or regular dash (-), we'll use regular dash to match filename
    for para in doc.paragraphs[:5]:  # Check first few paragraphs
        text = para.text
        # Check if this looks like the title paragraph (contains placeholders)
        if "<Full Name>" in text or "<Year>" in text or ("Brag Document" in text and ("<" in text or "Full Name" in text)):
            # Replace placeholders - title should match filename format exactly
            # Filename format: "Brag Document - <Full Name> (<Year>)"
            # Replace any dash variant (em dash, en dash, regular dash) with regular dash
            new_text = f"Brag Document - {full_name} ({year})"
            # Clear and set new text
            para.clear()
            para.add_run(new_text)
            break
    
    # Remove example entries (paragraphs containing "example" in text)
    paragraphs_to_remove = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        # Remove paragraphs that are example entries
        if text and ("example" in text or 
                   (text.startswith("•") and "example" in text) or
                   (text.startswith("-") and "example" in text)):
            paragraphs_to_remove.append(i)
    
    # Remove in reverse order to maintain indices
    for i in reversed(paragraphs_to_remove):
        para_element = doc.paragraphs[i]._element
        para_element.getparent().remove(para_element)
    
    doc.save(str(doc_path))


def add_entry_to_index(
    index_path: Path,
    entry_id: str,
    section_path: str,
    text: str,
    paragraph_index: int
) -> None:
    """Add an entry to the index file.
    
    Args:
        index_path: Path to the index file.
        entry_id: Unique identifier for the entry.
        section_path: Section path where the entry was added.
        text: Text content of the entry.
        paragraph_index: Paragraph index in the document.
    """
    index_data = load_index(index_path)
    
    # Ensure section exists in index
    if section_path not in index_data["sections"]:
        index_data["sections"][section_path] = {
            "created_at": datetime.now().isoformat(),
            "entry_ids": []
        }
    
    # Add entry to index
    entry_data = {
        "entry_id": entry_id,
        "section_path": section_path,
        "text": text,
        "paragraph_index": paragraph_index,
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat()
    }
    
    index_data["entries"][entry_id] = entry_data
    index_data["sections"][section_path]["entry_ids"].append(entry_id)
    
    save_index(index_path, index_data)
